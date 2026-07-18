"""Evaluation framework for PII detection.

Computes Precision, Recall, and F1 score per detector type
using exact match for structured PII and containment match for NER-based PII.
"""

import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))

from docx import Document
from pii_redactor.detectors import (
    detect_emails_in_text,
    detect_phones_in_text,
    detect_persons_in_text,
    detect_companies_in_text,
    detect_addresses_in_text,
    detect_ipv4_in_text,
    detect_ipv6_in_text,
    detect_dates_in_text,
    detect_credit_cards_in_text,
    detect_national_ids_in_text,
)


DOCUMENTS_DIR = Path(__file__).parent / "documents"


def load_ground_truth(blind: bool = False) -> dict:
    """Load ground truth labels from JSON."""
    suffix = "blind_" if blind else ""
    gt_path = DOCUMENTS_DIR / f"{suffix}ground_truth.json"
    with open(gt_path) as f:
        return json.load(f)


def extract_full_text(docx_path: Path) -> str:
    """Extract all text from a DOCX document."""
    doc = Document(str(docx_path))
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        texts.append(para.text)
    return "\n".join(texts)


def get_detected_entities(docx_path: Path) -> dict[str, list[str]]:
    """Run detectors on document text and extract detected entities by type."""
    full_text = extract_full_text(docx_path)

    detectors = {
        "email": detect_emails_in_text,
        "phone": detect_phones_in_text,
        "person": detect_persons_in_text,
        "company": detect_companies_in_text,
        "address": detect_addresses_in_text,
        "ipv4": detect_ipv4_in_text,
        "ipv6": detect_ipv6_in_text,
        "date": detect_dates_in_text,
        "credit_card": detect_credit_cards_in_text,
        "national_id": detect_national_ids_in_text,
    }

    entities_by_type: dict[str, list[str]] = {}
    for pii_type, detect_fn in detectors.items():
        detections = detect_fn(full_text)
        entities_by_type[pii_type] = [d.text for d in detections]

    return entities_by_type


def _normalize(text: str) -> str:
    """Normalize text for comparison."""
    return text.strip().lower()


def compute_metrics(
    ground_truth: list[dict],
    detected: dict[str, list[str]],
    match_mode: dict[str, str],
) -> dict:
    """Compute Precision, Recall, F1 per entity type."""
    results = {}

    gt_by_type: dict[str, list[str]] = {}
    for item in ground_truth:
        t = item["type"]
        if t not in gt_by_type:
            gt_by_type[t] = []
        gt_by_type[t].append(item["text"])

    for entity_type, gt_texts in gt_by_type.items():
        mode = match_mode.get(entity_type, "exact")
        detected_texts = detected.get(entity_type, [])
        detected_norm = [_normalize(t) for t in detected_texts]

        gt_detected = 0
        for gt_text in gt_texts:
            gt_norm = _normalize(gt_text)
            if mode == "contain":
                found = any(
                    gt_norm in dn or dn in gt_norm
                    for dn in detected_norm
                )
            else:
                found = gt_norm in detected_norm
            if found:
                gt_detected += 1

        pred_correct = 0
        for dn in detected_norm:
            if mode == "contain":
                found = any(
                    _normalize(gt) in dn or dn in _normalize(gt)
                    for gt in gt_texts
                )
            else:
                found = dn in [_normalize(gt) for gt in gt_texts]
            if found:
                pred_correct += 1

        precision = pred_correct / len(detected_texts) if detected_texts else 0.0
        recall = gt_detected / len(gt_texts) if gt_texts else 0.0
        f1 = (
            2 * precision * recall / (precision + recall)
            if (precision + recall) > 0
            else 0.0
        )

        results[entity_type] = {
            "true_positives": gt_detected,
            "pred_correct": pred_correct,
            "predicted": len(detected_texts),
            "expected": len(gt_texts),
            "precision": round(precision, 4),
            "recall": round(recall, 4),
            "f1": round(f1, 4),
        }

    return results


def evaluate_all(blind: bool = False) -> dict:
    """Evaluate all synthetic documents against ground truth."""
    metadata = load_ground_truth(blind=blind)
    match_mode = {
        "person": "contain",
        "company": "contain",
        "address": "contain",
        "email": "exact",
        "phone": "exact",
        "date": "exact",
        "ipv4": "exact",
        "ipv6": "exact",
        "credit_card": "exact",
        "national_id": "exact",
    }

    all_results = {}
    for doc_id, doc_info in metadata.items():
        docx_path = DOCUMENTS_DIR / doc_info["filename"]
        if not docx_path.exists():
            print(f"  Skipping {doc_info['filename']} (not found)")
            continue

        detected = get_detected_entities(docx_path)
        metrics = compute_metrics(doc_info["ground_truth"], detected, match_mode)
        all_results[doc_id] = metrics

    return all_results


def print_results(results: dict) -> None:
    """Print evaluation results in a readable format."""
    for doc_id, metrics in results.items():
        print(f"\n--- {doc_id} ---")
        for entity_type, m in metrics.items():
            print(
                f"  {entity_type:15s}: "
                f"P={m['precision']:.2f}  R={m['recall']:.2f}  F1={m['f1']:.2f}  "
                f"({m['true_positives']}/{m['predicted']} pred, "
                f"{m['expected']} expected)"
            )


def compute_overall(results: dict) -> dict:
    """Compute aggregate metrics across all documents."""
    totals: dict[str, dict] = {}

    for doc_metrics in results.values():
        for entity_type, m in doc_metrics.items():
            if entity_type not in totals:
                totals[entity_type] = {"tp": 0, "pred_correct": 0, "pred": 0, "exp": 0}
            totals[entity_type]["tp"] += m["true_positives"]
            totals[entity_type]["pred_correct"] += m.get("pred_correct", m["true_positives"])
            totals[entity_type]["pred"] += m["predicted"]
            totals[entity_type]["exp"] += m["expected"]

    overall = {}
    for entity_type, t in totals.items():
        p = t["pred_correct"] / t["pred"] if t["pred"] > 0 else 0.0
        r = t["tp"] / t["exp"] if t["exp"] > 0 else 0.0
        f1 = 2 * p * r / (p + r) if (p + r) > 0 else 0.0
        overall[entity_type] = {
            "precision": round(p, 4),
            "recall": round(r, 4),
            "f1": round(f1, 4),
            "true_positives": t["tp"],
            "predicted": t["pred"],
            "expected": t["exp"],
        }

    return overall


def main():
    """Run evaluation and print results."""
    print("Running PII detection on synthetic documents...\n")
    results = evaluate_all(blind=False)

    print("\n=== Per-Document Results ===")
    print_results(results)

    print("\n=== Overall Metrics ===")
    overall = compute_overall(results)
    for entity_type, m in overall.items():
        print(
            f"  {entity_type:15s}: "
            f"P={m['precision']:.2f}  R={m['recall']:.2f}  F1={m['f1']:.2f}  "
            f"({m['true_positives']}/{m['predicted']} pred, "
            f"{m['expected']} expected)"
        )

    if overall:
        avg_p = sum(m["precision"] for m in overall.values()) / len(overall)
        avg_r = sum(m["recall"] for m in overall.values()) / len(overall)
        avg_f1 = sum(m["f1"] for m in overall.values()) / len(overall)
        print(f"\n  {'MACRO AVG':15s}: P={avg_p:.2f}  R={avg_r:.2f}  F1={avg_f1:.2f}")

    output_path = DOCUMENTS_DIR / "evaluation_results.json"
    with open(output_path, "w") as f:
        json.dump({"per_document": results, "overall": overall}, f, indent=2)
    print(f"\nResults saved to: {output_path}")

    print("\n" + "=" * 60)
    print("BLIND VALIDATION")
    print("=" * 60)

    print("\nRunning PII detection on blind documents...\n")
    blind_results = evaluate_all(blind=True)

    print("\n=== Blind Per-Document Results ===")
    print_results(blind_results)

    print("\n=== Blind Overall Metrics ===")
    blind_overall = compute_overall(blind_results)
    for entity_type, m in blind_overall.items():
        print(
            f"  {entity_type:15s}: "
            f"P={m['precision']:.2f}  R={m['recall']:.2f}  F1={m['f1']:.2f}  "
            f"({m['true_positives']}/{m['predicted']} pred, "
            f"{m['expected']} expected)"
        )

    if blind_overall:
        avg_p = sum(m["precision"] for m in blind_overall.values()) / len(blind_overall)
        avg_r = sum(m["recall"] for m in blind_overall.values()) / len(blind_overall)
        avg_f1 = sum(m["f1"] for m in blind_overall.values()) / len(blind_overall)
        print(f"\n  {'MACRO AVG':15s}: P={avg_p:.2f}  R={avg_r:.2f}  F1={avg_f1:.2f}")

    blind_output_path = DOCUMENTS_DIR / "blind_evaluation_results.json"
    with open(blind_output_path, "w") as f:
        json.dump({"per_document": blind_results, "overall": blind_overall}, f, indent=2)
    print(f"\nBlind results saved to: {blind_output_path}")

    return overall, blind_overall


if __name__ == "__main__":
    main()
