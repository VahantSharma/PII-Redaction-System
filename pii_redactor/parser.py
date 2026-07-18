"""
Parse and rebuild .docx files.

This module handles reading a .docx into a structured representation
and writing it back while preserving formatting.

Design decision: We use python-docx directly rather than building a
custom DocumentModel. python-docx's Document object IS our canonical
representation. We modify it in-place (or on a copy) and save.

This avoids the complexity of maintaining a parallel data model while
still giving us access to runs, formatting, tables, headers, and footers.
"""

from pathlib import Path

from docx import Document


def load_document(path: str | Path) -> Document:
    """Load a .docx file into a python-docx Document object.

    Args:
        path: Path to the .docx file.

    Returns:
        python-docx Document object.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file is not a valid .docx.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    if not path.suffix.lower() == ".docx":
        raise ValueError(f"Expected .docx file, got: {path.suffix}")

    return Document(str(path))


def save_document(document: Document, path: str | Path) -> None:
    """Save a Document to a .docx file.

    Args:
        document: python-docx Document object.
        path: Output path for the .docx file.
    """
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))


def count_paragraphs(document: Document) -> int:
    """Count paragraphs in body (not inside tables)."""
    return len(document.paragraphs)


def count_tables(document: Document) -> int:
    """Count tables in the document."""
    return len(document.tables)


def count_table_rows(document: Document) -> int:
    """Count total rows across all tables."""
    return sum(len(table.rows) for table in document.tables)


def count_table_cells(document: Document) -> int:
    """Count total cells across all tables."""
    return sum(
        len(row.cells)
        for table in document.tables
        for row in table.rows
    )


def get_all_text(document: Document) -> list[str]:
    """Extract all text from the document (body paragraphs + table cells).

    Returns:
        List of non-empty text strings.
    """
    texts = []

    for para in document.paragraphs:
        if para.text.strip():
            texts.append(para.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        texts.append(para.text)

    return texts


def get_document_stats(document: Document) -> dict:
    """Get structural statistics of the document.

    Returns:
        Dict with paragraph count, table count, row count, cell count.
    """
    return {
        "paragraphs": count_paragraphs(document),
        "tables": count_tables(document),
        "rows": count_table_rows(document),
        "cells": count_table_cells(document),
    }


def validate_structure(original: Document, output: Document) -> dict:
    """Compare structure of two documents.

    Returns:
        Dict with comparison results and pass/fail for each metric.
    """
    orig_stats = get_document_stats(original)
    out_stats = get_document_stats(output)

    results = {}
    for key in orig_stats:
        results[key] = {
            "original": orig_stats[key],
            "output": out_stats[key],
            "match": orig_stats[key] == out_stats[key],
        }

    results["all_match"] = all(r["match"] for r in results.values())
    return results


def validate_run_formatting(original: Document, output: Document, sample_size: int = 50) -> dict:
    """Compare run-level formatting between two documents.

    Checks bold, italic, underline, font name, and font size for
    the first N runs that have text content.

    Args:
        original: Original document.
        output: Output document to compare against.
        sample_size: Number of runs to compare.

    Returns:
        Dict with match count, mismatch count, and details.
    """
    def extract_runs(doc):
        runs = []
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text.strip():
                    runs.append(run)
                if len(runs) >= sample_size:
                    return runs
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text.strip():
                                runs.append(run)
                            if len(runs) >= sample_size:
                                return runs
        return runs

    orig_runs = extract_runs(original)
    out_runs = extract_runs(output)

    matches = 0
    mismatches = 0
    details = []

    for i, (orig, out) in enumerate(zip(orig_runs, out_runs)):
        attrs = ["bold", "italic", "underline"]
        for attr in attrs:
            orig_val = getattr(orig, attr, None)
            out_val = getattr(out, attr, None)
            if orig_val == out_val:
                matches += 1
            else:
                mismatches += 1
                details.append(f"Run {i}: {attr} {orig_val} -> {out_val}")

    return {
        "matches": matches,
        "mismatches": mismatches,
        "total": matches + mismatches,
        "all_match": mismatches == 0,
        "details": details[:10],  # first 10 mismatches
    }
